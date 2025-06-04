from fastapi import APIRouter, UploadFile, File, Form
from fastapi.responses import FileResponse
from docx import Document
from docx.shared import RGBColor
from datetime import datetime
import json
import os

router = APIRouter(prefix="/thisismagic/contracts", tags=["contracts"])

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "templates")
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "generated")

@router.post("/generate")
async def generate_contract(
    template_name: str = Form(...),
    data: str = Form(...)
):
    # Convertimos los datos a diccionario
    campos = json.loads(data)

    # Construimos el path del template con extensión .docx
    template_path = os.path.join(TEMPLATES_DIR, f"{template_name}.docx")
    if not os.path.exists(template_path):
        return {"error": f"No se encontró el template {template_name}.docx"}

    # Cargar documento y reemplazar <<variables>> (en párrafos)
    doc = Document(template_path)
    for p in doc.paragraphs:
        full_text = ''.join(run.text for run in p.runs)
        for key, value in campos.items():
            placeholder = f"<<{key}>>"
            if placeholder in full_text:
                full_text = full_text.replace(placeholder, str(value))
                for i in range(len(p.runs)):
                    p.runs[i].text = ''
                if p.runs:
                    r = p.runs[0]
                    r.text = full_text
                    

    # Reemplazar también en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full_text = ''.join(run.text for run in p.runs)
                    for key, value in campos.items():
                        placeholder = f"<<{key}>>"
                        if placeholder in full_text:
                            full_text = full_text.replace(placeholder, str(value))
                            for i in range(len(p.runs)):
                                p.runs[i].text = ''
                            if p.runs:
                                r = p.runs[0]
                                r.text = full_text
                                

    # Guardar el resultado temporal
    output_path = os.path.join(OUTPUT_DIR, f"{template_name}_filled.docx")
    doc.save(output_path)

    today = datetime.now()

    return FileResponse(output_path, filename=f"{template_name}_generado_{today}.docx")
